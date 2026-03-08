from __future__ import annotations

from pathlib import Path
from typing import List

from docx import Document


def load_document(docx_path: Path) -> Document:
    """Load a DOCX file from disk."""
    return Document(str(docx_path))


def extract_tables(doc: Document) -> List[List[List[str]]]:
    """Extract tables from DOCX as 3D list: table -> row -> cell text."""
    all_tables: List[List[List[str]]] = []

    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(cells)
        all_tables.append(rows)

    return all_tables


def extract_readable_text(doc: Document) -> str:
    """Extract readable content from paragraphs and tables for LLM processing."""
    chunks: List[str] = []

    paragraph_text = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    if paragraph_text:
        chunks.append("\n".join(paragraph_text))

    for index, table in enumerate(doc.tables, start=1):
        chunks.append(f"\n[DOCX_TABLE_{index}]\n")
        for row in table.rows:
            row_text = " | ".join(cell.text.strip().replace("\n", " ") for cell in row.cells)
            chunks.append(row_text)

    return "\n".join(chunks)
